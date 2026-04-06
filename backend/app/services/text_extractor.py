"""
Text extraction service for PDF and DOCX documents.

Supports:
- PDF via pdfplumber (preferred) with PyPDF2 fallback
- DOCX via python-docx including table content
- Plain text files with UTF-8 decoding

All extractors return a single string with page/section separators.
"""

import logging
import io
from typing import Optional

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# PDF extraction
# ---------------------------------------------------------------------------

def extract_text_from_pdf(content: bytes) -> str:
    """
    Extract text from a PDF file.

    Tries pdfplumber first (better layout preservation), falls back to PyPDF2.
    Each page is prefixed with a [Page N] header for traceability.

    Args:
        content: Raw bytes of the PDF file.

    Returns:
        Extracted text as a single string, or empty string on failure.
    """
    try:
        import pdfplumber
        text_parts = []
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            for page_num, page in enumerate(pdf.pages, 1):
                page_text = page.extract_text()
                if page_text and page_text.strip():
                    text_parts.append(f"[Page {page_num}]\n{page_text.strip()}")
        extracted = "\n\n".join(text_parts)
        logger.debug(f"pdfplumber extracted {len(extracted)} chars from {len(text_parts)} pages")
        return extracted

    except ImportError:
        logger.warning("pdfplumber not installed — falling back to PyPDF2")
        return _extract_pdf_pypdf2(content)
    except Exception as e:
        logger.error(f"pdfplumber extraction failed: {e} — attempting PyPDF2 fallback")
        return _extract_pdf_pypdf2(content)


def _extract_pdf_pypdf2(content: bytes) -> str:
    """
    Fallback PDF text extraction using PyPDF2.

    Args:
        content: Raw bytes of the PDF file.

    Returns:
        Extracted text, or empty string on failure.
    """
    try:
        import PyPDF2
        reader = PyPDF2.PdfReader(io.BytesIO(content))
        text_parts = []
        for page_num, page in enumerate(reader.pages, 1):
            page_text = page.extract_text()
            if page_text and page_text.strip():
                text_parts.append(f"[Page {page_num}]\n{page_text.strip()}")
        extracted = "\n\n".join(text_parts)
        logger.debug(f"PyPDF2 extracted {len(extracted)} chars from {len(text_parts)} pages")
        return extracted
    except ImportError:
        logger.error("Neither pdfplumber nor PyPDF2 is installed. Cannot extract PDF text.")
        return ""
    except Exception as e:
        logger.error(f"PyPDF2 extraction failed: {e}")
        return ""


# ---------------------------------------------------------------------------
# DOCX extraction
# ---------------------------------------------------------------------------

def extract_text_from_docx(content: bytes) -> str:
    """
    Extract text from a DOCX file using python-docx.

    Extracts:
    - All non-empty paragraphs in document order
    - All table cell content (formatted as pipe-delimited rows)

    Args:
        content: Raw bytes of the DOCX file.

    Returns:
        Extracted text as a single string, or empty string on failure.
    """
    try:
        import docx
        doc = docx.Document(io.BytesIO(content))
        text_parts = []

        # Extract paragraph text
        for para in doc.paragraphs:
            stripped = para.text.strip()
            if stripped:
                text_parts.append(stripped)

        # Extract table content
        for table_idx, table in enumerate(doc.tables, 1):
            table_parts = [f"[Table {table_idx}]"]
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                row_text = " | ".join(c for c in cells if c)
                if row_text:
                    table_parts.append(row_text)
            if len(table_parts) > 1:  # Only add if the table had content
                text_parts.append("\n".join(table_parts))

        extracted = "\n\n".join(text_parts)
        logger.debug(f"python-docx extracted {len(extracted)} chars")
        return extracted

    except ImportError:
        logger.error("python-docx is not installed. Cannot extract DOCX text.")
        return ""
    except Exception as e:
        logger.error(f"DOCX extraction failed: {e}")
        return ""


# ---------------------------------------------------------------------------
# Plain text extraction
# ---------------------------------------------------------------------------

def extract_text_from_txt(content: bytes) -> str:
    """
    Decode a plain text file.

    Attempts UTF-8 first, falls back to latin-1.

    Args:
        content: Raw bytes of the text file.

    Returns:
        Decoded text string.
    """
    for encoding in ("utf-8", "utf-8-sig", "latin-1", "cp1252"):
        try:
            decoded = content.decode(encoding)
            logger.debug(f"Text file decoded with {encoding} ({len(decoded)} chars)")
            return decoded
        except UnicodeDecodeError:
            continue
    logger.error("Failed to decode text file with any supported encoding")
    return content.decode("utf-8", errors="replace")


# ---------------------------------------------------------------------------
# Unified extraction entry point
# ---------------------------------------------------------------------------

def extract_text(content: bytes, file_extension: str) -> str:
    """
    Extract text from a document based on its file extension.

    Args:
        content: Raw bytes of the file.
        file_extension: File extension including dot, e.g. '.pdf', '.docx', '.txt'.

    Returns:
        Extracted text string (may be empty if extraction failed or file has no text).
    """
    ext = file_extension.lower().strip()

    if ext == ".pdf":
        return extract_text_from_pdf(content)
    elif ext in (".docx", ".doc"):
        return extract_text_from_docx(content)
    elif ext == ".txt":
        return extract_text_from_txt(content)
    else:
        logger.warning(f"Unsupported file extension for text extraction: '{ext}'")
        return ""


def truncate_text(text: str, max_chars: int = 50000) -> str:
    """
    Truncate text to a maximum character count suitable for API calls.

    Appends a notice when truncation occurs so the AI knows the document
    was not fully analysed.

    Args:
        text: The full extracted text.
        max_chars: Maximum characters to retain.

    Returns:
        Original text if within limit, or truncated text with a notice.
    """
    if not text or len(text) <= max_chars:
        return text
    truncated = text[:max_chars]
    logger.debug(f"Text truncated from {len(text)} to {max_chars} chars")
    return truncated + "\n\n[... document truncated for analysis — remaining content not shown ...]"
